import os
import uuid
import pandas as pd
from docx import Document
import yagmail

OUTPUT_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'outputs', 'appointment_letters'))
os.makedirs(OUTPUT_DIR, exist_ok=True)

def load_and_validate_data(file_path):
    required_columns = {'Name', 'Email', 'DOB', 'Phone'}
    df = pd.read_excel(file_path)
    missing = required_columns - set(df.columns)
    if missing:
        raise ValueError(f"Excel file is missing columns: {missing}")
    return df

def generate_gid():
    return "LTC-" + str(uuid.uuid4())[:7]

def create_appointment_letter(name, gid, output_folder=OUTPUT_DIR):
    doc = Document()
    doc.add_heading('Appointment Letter', 0)
    doc.add_paragraph(f"Dear {name},")
    doc.add_paragraph("We are pleased to offer you a position at LTC.")
    doc.add_paragraph(f"Your Employee ID: {gid}")
    doc.add_paragraph("Please report to the HR department for onboarding.")
    filepath = os.path.join(output_folder, f"{gid}_Appointment_Letter.docx")
    doc.save(filepath)
    return filepath

def send_onboarding_email(to_email, name, gid, attachment_path, yag):
    subject = "Welcome to LTC - Your Appointment Letter"
    body = f"""Dear {name},

Welcome to LTC!

Your Employee ID is {gid}.

Please find your appointment letter attached.

Classification: Highly Confidential 

I’m delighted to make you an offer for the role of Associate Engineer, Lloyds Technology Centre. 

Congratulations and we can’t wait for you to join us. Before you do, there are just a few things we need from you.  Please review and accept your conditional offer letter within three working days. Upon offer acceptance we will need to commence the vetting process which can take up to four weeks.  Once we initiate your vetting process you will be receiving an email from our vetting partner (HireRight) with instructions on how to complete background checks.

Please note the Relocation Benefit may be utilised towards:

Booking temporary accommodation (hotel/serviced apartment) at your Location of Joining until you find permanent accommodation
Movement of personal goods and vehicles
Expenses incurred on cab/taxi fares
Brokerage fees and society move in charges (as applicable)
Any other expenses incurred during relocation
 

Your one-time travel expenses for (bus/ air/train tickets) from your current location to your Location of Joining would be reimbursed to you at actuals post joining.

We look forward to welcoming you to Lloyds Technology Centre!

Best regards,
HR Team
"""
    yag.send(to=to_email, subject=subject, contents=body, attachments=attachment_path)

def process_onboarding(excel_path, sender_email, sender_password,
                       email_host='smtp.gmail.com', email_port=587,
                       smtp_ssl=False, smtp_starttls=True):

    df = load_and_validate_data(excel_path)

    df['GID'] = df.apply(lambda _: generate_gid(), axis=1)
    df['LetterPath'] = df.apply(lambda row: create_appointment_letter(row['Name'], row['GID']), axis=1)

    yag = yagmail.SMTP(
        user=sender_email,
        password=sender_password,
        host=email_host,
        port=email_port,
        smtp_ssl=smtp_ssl,
        smtp_starttls=smtp_starttls
    )

    sent_emails = []
    failed_emails = []

    for _, row in df.iterrows():
        try:
            send_onboarding_email(row['Email'], row['Name'], row['GID'], row['LetterPath'], yag)
            sent_emails.append(row['Email'])
        except Exception as e:
            failed_emails.append((row['Email'], str(e)))

    return sent_emails, failed_emails

import os
import uuid
import pandas as pd
from docx import Document
import yagmail

SENDGRID_SMTP_SERVER = "smtp.sendgrid.net"
SENDGRID_SMTP_PORT = 587
SENDER_EMAIL = "HRLloydsbanking.com"
SENDGRID_API_KEY = "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"

DATA_FILE = os.path.join('..', 'data', 'random_user_data.xlsx')
OUTPUT_DIR = os.path.join('..', 'outputs', 'appointment_letters')
os.makedirs(OUTPUT_DIR, exist_ok=True)

def load_and_validate_data(file_path):
    required_columns = {'Name', 'Email', 'DOB', 'Phone'}
    df = pd.read_excel(file_path)
    missing = required_columns - set(df.columns)
    if missing:
        raise ValueError(f"Excel missing columns: {missing}")
    if df['Email'].isnull().any():
        print("Warning: Some rows are missing emails.")
    return df

def generate_gid():
    return "LTC-" + str(uuid.uuid4())[:7]

def create_appointment_letter(name, gid, output_folder):
    doc = Document()
    doc.add_heading('Appointment Letter', level=0)
    doc.add_paragraph(f"Dear {name},")
    doc.add_paragraph("We are pleased to offer you a position at LTC.")
    doc.add_paragraph(f"Your Employee ID: {gid}")
    doc.add_paragraph("Please report to the HR department for onboarding.")
    letter_path = os.path.join(output_folder, f"{gid}_Appointment_Letter.docx")
    doc.save(letter_path)
    print(f"Created letter at {letter_path}")
    return letter_path

def send_onboarding_email(recipient_email, name, gid, attachment, yag):
    subject = "Welcome to LTC — Your Appointment Letter"
    body = f"""Dear {name},

Welcome to LTC!

Your Employee ID is {gid}.

Please find your appointment letter attached.

Classification: Highly Confidential 

I’m delighted to make you an offer for the role of Associate Engineer, Lloyds Technology Centre. 

Congratulations and we can’t wait for you to join us. Before you do, there are just a few things we need from you.  Please review and accept your conditional offer letter within three working days. Upon offer acceptance we will need to commence the vetting process which can take up to four weeks.  Once we initiate your vetting process you will be receiving an email from our vetting partner (HireRight) with instructions on how to complete background checks.

Please note the Relocation Benefit may be utilised towards:

Booking temporary accommodation (hotel/serviced apartment) at your Location of Joining until you find permanent accommodation
Movement of personal goods and vehicles
Expenses incurred on cab/taxi fares
Brokerage fees and society move in charges (as applicable)
Any other expenses incurred during relocation
 

Your one-time travel expenses for (bus/ air/train tickets) from your current location to your Location of Joining would be reimbursed to you at actuals post joining.

We look forward to welcoming you to Lloyds Technology Centre!


Best regards,
HR Team
"""
    try:
        yag.send(
            to=recipient_email,
            subject=subject,
            contents=body,
            attachments=attachment
        )
        print(f"Email sent to {recipient_email}")
    except Exception as e:
        print(f"Failed to send email to {recipient_email}: {e}")

def main():
    print("Loading employee data...")
    df = load_and_validate_data(DATA_FILE)

    print("Generating GIDs and appointment letters...")
    df['GID'] = df.apply(lambda _: generate_gid(), axis=1)
    df['LetterPath'] = df.apply(lambda row: create_appointment_letter(row['Name'], row['GID'], OUTPUT_DIR), axis=1)

    print("Connecting to SendGrid SMTP...")
    yag = yagmail.SMTP(
        user=SENDER_EMAIL,
        password=SENDGRID_API_KEY,
        host=SENDGRID_SMTP_SERVER,
        port=SENDGRID_SMTP_PORT
    )

    print("Sending onboarding emails with appointment letters attached...")
    for _, row in df.iterrows():
        send_onboarding_email(row['Email'], row['Name'], row['GID'], row['LetterPath'], yag)

    print("All emails sent successfully.")

if __name__ == "__main__":
    main()
